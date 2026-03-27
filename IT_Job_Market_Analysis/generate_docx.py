import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_docx(input_md, output_docx):
    doc = Document()
    
    # Title
    title = doc.add_heading("BÁO CÁO DỰ ÁN BUSINESS INTELLIGENCE", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if not os.path.exists(input_md):
        print(f"Error: {input_md} not found.")
        return

    with open(input_md, "r", encoding="utf-8") as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith("# "):
            h = doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            h = doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            h = doc.add_heading(line[4:], level=3)
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        else:
            p = doc.add_paragraph(line.replace("**", ""))

    doc.save(output_docx)
    print(f"Successfully generated {output_docx}")

if __name__ == "__main__":
    generate_docx("Final_Project_Report.md", "Business_Sentiment_Analysis_Report.docx")
